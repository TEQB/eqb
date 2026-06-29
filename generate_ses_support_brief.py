from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "AWS_SES_Email_Response_Summary.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Aptos", size=10.5, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    run = p.add_run(text)
    set_font(run)
    return p


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + ": ")
    set_font(r1, bold=True, color="7A1030")
    r2 = p.add_run(value)
    set_font(r2)
    return p


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ("Title", 22, "111827"),
        ("Heading 1", 15, "7A1030"),
        ("Heading 2", 11.5, "7A1030"),
    ]:
        s = styles[name]
        s.font.name = "Aptos"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)


def main():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("AWS SES support case: email-sending summary")
    set_font(r, size=22, bold=True, color="111827")

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    r = sub.add_run("Evidence extracted from the PQB codebase for the SES production-access request.")
    set_font(r, size=10.5, color="4B5563")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Scope: ")
    set_font(r, bold=True, color="7A1030")
    r = p.add_run("backend route handlers, auth flows, notification handlers, configuration, and cron/webhook/bounce checks.")
    set_font(r)

    doc.add_heading("1. What triggers outbound email", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("The app does not have a marketing/newsletter mailer. Every email path I found is a direct response to a specific user or admin action.")
    set_font(r)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [1.55, 2.25, 1.55, 2.5]
    hdr = table.rows[0].cells
    headers = ["Trigger", "Who receives it", "What it sends", "Evidence"]
    for i, h in enumerate(headers):
        hdr[i].width = Inches(widths[i])
        hdr[i].text = h
        set_cell_shading(hdr[i], "E9D9DD")
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, bold=True, color="111827")
    set_repeat_table_header(table.rows[0])

    rows = [
        [
            "Student registration / OTP request",
            "The registering user",
            "A 6-digit sign-in / verification code",
            "src/app/api/auth/send-otp/route.ts and src/components/auth/RegisterForm.tsx; resend can also be triggered from src/components/auth/OtpForm.tsx",
        ],
        [
            "Student password reset",
            "The student who requested reset",
            "A password-reset link generated from Supabase Admin",
            "src/app/api/auth/forgot-password/route.ts and src/components/auth/LoginForm.tsx",
        ],
        [
            "Admin password reset",
            "The super admin who requested reset",
            "A password-reset link generated from Supabase Admin",
            "src/app/api/auth/forgot-password/route.ts and src/components/admin/AdminLoginForm.tsx",
        ],
        [
            "Admin invitation",
            "The invited admin email address",
            "Invitation email with recovery link or temporary password",
            "src/app/api/admin/route.ts case \"invite-admin\" and src/components/admin/AdminManagement.tsx",
        ],
        [
            "Student feedback submission",
            "All current super-admin recipients",
            "A notification email containing the student name, programme, and feedback text",
            "src/app/api/feedback/route.ts and src/lib/admin-recipients.ts",
        ],
    ]

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_font(run)

    doc.add_heading("2. Volume estimate", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("I could not find real usage counts in the repository, so I cannot derive a trustworthy emails/day or emails/month number from code alone.")
    set_font(r)
    add_bullet(doc, "The admin stats endpoint computes counts such as total students, total questions, and uploads today, but the repo does not contain live values. See src/app/api/admin/route.ts case \"stats\".")
    add_bullet(doc, "No analytics table, dashboard export, or usage log in the repo provides current registration, reset, or feedback message volume.")
    add_bullet(doc, "If you need a number for AWS, it will need to come from production logs or database analytics, not from the checked-in code.")

    doc.add_heading("3. Recipient list management", level=1)
    add_bullet(doc, "There is no newsletter list, bulk marketing mailer, or scheduled batch email job in the repo.")
    add_bullet(doc, "The only multi-recipient mail path is feedback, which looks up current super-admin emails from profiles/auth.users and sends one notification email with BCC to the other admins.")
    add_bullet(doc, "Admin invites are one-at-a-time and are initiated manually from the admin UI.")

    doc.add_heading("4. Bounce / complaint / unsubscribe handling", level=1)
    add_bullet(doc, "I found no bounce webhook, complaint webhook, SNS handler, suppression-list logic, or unsubscribe link handling in src/ or supabase/.")
    add_bullet(doc, "I also found no SES-specific webhook endpoint in the codebase.")
    add_bullet(doc, "The only related repo-level email config is a commented Supabase SMTP block in supabase/config.toml, not active bounce processing.")

    doc.add_heading("5. Sender identity and config", level=1)
    add_kv(doc, "Sender address used in custom emails", "EQB <noreply@devalyze.space>")
    add_kv(doc, "Configured domain for student email gating", "@tech-u.edu.ng")
    add_kv(doc, "Email provider currently used by custom code", "Resend API via https://api.resend.com/emails")
    add_kv(doc, "Required env vars related to email", "RESEND_API_KEY and NEXT_PUBLIC_UNIVERSITY_EMAIL_DOMAIN")
    add_kv(doc, "Supabase SMTP status", "Present only as commented-out example config in supabase/config.toml; not active in the checked-in config")

    doc.add_heading("6. Actual template content found in code", level=1)
    p = doc.add_paragraph()
    r = p.add_run("The repository uses inline HTML strings instead of separate template files. Below are short excerpts and the exact source files where each full template lives.")
    set_font(r)

    templates = [
        (
            "OTP email",
            "Subject: \"Your EQB sign-in code\"",
            "HTML includes: h2 \"Your sign-in code\", an expiration note, and the 6-digit code in a centered styled block.",
            "src/app/api/auth/send-otp/route.ts",
        ),
        (
            "Password reset",
            "Subject: \"Reset your EQB password\" / \"Reset your EQB admin password\"",
            "HTML includes: h2 \"Password Reset\", a reset button linking to the generated recovery link, and a 24-hour expiry note.",
            "src/app/api/auth/forgot-password/route.ts",
        ),
        (
            "Admin invite",
            "Subject: \"You've been invited as an admin on EQB\"",
            "HTML includes: h2 \"Admin Invitation\", a setup-password button when a recovery link is present, or a temporary password fallback.",
            "src/app/api/admin/route.ts case \"invite-admin\"",
        ),
        (
            "Feedback notification",
            "Subject: \"New EQB student feedback\"",
            "HTML includes: student name, programme, submitted timestamp, and the feedback body.",
            "src/app/api/feedback/route.ts",
        ),
    ]

    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.autofit = False
    widths2 = [1.15, 2.8, 2.95]
    hdr = t.rows[0].cells
    for i, h in enumerate(["Template", "Exact inline subject / excerpt", "Source file"]):
        hdr[i].width = Inches(widths2[i])
        hdr[i].text = h
        set_cell_shading(hdr[i], "E9D9DD")
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, bold=True, color="111827")
    set_repeat_table_header(t.rows[0])

    for name, subject, excerpt, source in templates:
        cells = t.add_row().cells
        vals = [name, f"{subject}\n{excerpt}", source]
        for i, val in enumerate(vals):
            cells[i].width = Inches(widths2[i])
            cells[i].text = val
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_font(run)

    doc.add_heading("7. Bottom line for AWS", level=1)
    p = doc.add_paragraph()
    r = p.add_run("The app sends transactional email only, not bulk marketing email. Current custom sending is limited to OTP, password reset, admin invite, and feedback notification. There is no implemented bounce/complaint/unsubscribe workflow in the repo, and no live usage data to justify a numerical volume estimate from code alone.")
    set_font(r)

    doc.save(OUT)


if __name__ == "__main__":
    main()
